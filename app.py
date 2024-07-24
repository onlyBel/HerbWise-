from flask import Flask, render_template, request, send_file, session, redirect, url_for
import docx
import csv
import os
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime

app = Flask(__name__)
app.secret_key = "1111"
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///highlights.db'
db = SQLAlchemy(app)

class Document(db.Model):
    doc_id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    highlights = db.relationship('Highlight', backref='document', lazy=True)

class Highlight(db.Model):
    highlight_id = db.Column(db.Integer, primary_key=True)
    doc_id = db.Column(db.Integer, db.ForeignKey('document.doc_id'), nullable=False)
    color = db.Column(db.String(50), nullable=False)
    text = db.Column(db.Text, nullable=False)

def get_color_name(highlight_color):
    if isinstance(highlight_color, docx.enum.text.WD_COLOR_INDEX):
        return highlight_color.name
    else:
        return None

def extract_highlighted_text(docx_file):
    doc = docx.Document(docx_file)
    highlighted_data = {}

    for para in doc.paragraphs:
        for run in para.runs:
            highlight_color = run.font.highlight_color
            color_name = get_color_name(highlight_color)

            if color_name:
                highlighted_data.setdefault(color_name, []).append(run.text)

    return highlighted_data

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        docx_file = request.files['docxFile']
        if docx_file and docx_file.filename.endswith('.docx'):
            temp_path = os.path.join('uploads', docx_file.filename)
            docx_file.save(temp_path)

            highlighted_data = extract_highlighted_text(temp_path)
            os.remove(temp_path)

            new_doc = Document(filename=docx_file.filename)
            db.session.add(new_doc)
            db.session.commit()

            for color, texts in highlighted_data.items():
                for text in texts:
                    new_highlight = Highlight(doc_id=new_doc.doc_id, color=color, text=text)
                    db.session.add(new_highlight)
            db.session.commit()

            return redirect(url_for('index'))

    documents = Document.query.all()
    return render_template('index.html', documents=documents)

@app.route('/view/<int:doc_id>')
def view_highlights(doc_id):
    document = Document.query.get_or_404(doc_id)
    highlights = Highlight.query.filter_by(doc_id=doc_id).all()
    
    # Construct the highlighted_data dictionary
    highlighted_data = {}
    for highlight in highlights:
        highlighted_data.setdefault(highlight.color, []).append(highlight.text)
    
    return render_template('highlights.html', document=document, highlighted_data=highlighted_data)

@app.route('/download/<int:doc_id>')
def download_file(doc_id):
    document = Document.query.get_or_404(doc_id)
    highlights = Highlight.query.filter_by(doc_id=doc_id).all()

    output_file = os.path.join('outputs', f'highlighted_text_{doc_id}.csv')
    with open(output_file, "w", newline="", encoding="utf-8") as csv_file:
        fieldnames = ['Color', 'Highlighted Text']
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL, lineterminator='\n')

        writer.writeheader()

        for highlight in highlights:
            writer.writerow({'Color': highlight.color, 'Highlighted Text': highlight.text})

    return send_file(output_file, as_attachment=True)

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    os.makedirs('outputs', exist_ok=True)
    with app.app_context():
        db.create_all()
    app.run(debug=True)
